"""Gera o relatorio tecnico final (.docx) a partir de resultados.csv, dataset.json,
fase7_comparacao_grafo.json e ragas_*_medias.json, seguindo a estrutura da Secao 8
do roteiro do trabalho final."""
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PASTA = Path(__file__).resolve().parent
AVAL = PASTA.parent
GRAF = PASTA.parent / "graficos"

AZUL = RGBColor(0x1F, 0x38, 0x64)
CINZA = RGBColor(0x40, 0x40, 0x40)


def carregar_csv():
    with open(AVAL / "resultados.csv", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def carregar_json(nome):
    p = AVAL / nome
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for i in range(1, 4):
        h = doc.styles[f"Heading {i}"]
        h.font.color.rgb = AZUL
        h.font.name = "Calibri"


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n")
    run = p.add_run("A Jornada de Melhoria da Recuperação em RAG")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = AZUL
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Estudo de caso: Constituição da República Federativa do Brasil de 1988")
    run2.font.size = Pt(16)
    run2.font.color.rgb = CINZA

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(
        "Disciplina: RAG & CAG Aplicados a Direito e Segurança Pública\n"
        "Projeto base: Projeto Final (aula12/projeto_final) — API de Ingestão Inteligente + RAG\n"
        "Fonte de dados: Constituicao-Compilado.pdf (164 páginas, 250 artigos)\n"
    ).font.size = Pt(12)

    doc.add_paragraph()
    just = doc.add_paragraph()
    just.add_run("Por que esta fonte: ").bold = True
    just.add_run(
        "a Constituição é um texto jurídico longo, denso em artigos e incisos, com "
        "bastante remissão interna entre temas (direitos, poderes, tributos, ordem "
        "econômica). Isso cria um terreno rico para comparar estratégias de chunking, "
        "modelos de embedding, técnicas de recuperação (léxica, densa, híbrida, "
        "reranking) e uma técnica de grafo de conhecimento em perguntas multi-hop."
    )
    doc.add_page_break()


def add_heading(doc, texto, nivel=1):
    doc.add_heading(texto, level=nivel)


def add_metric_table(doc, linhas, colunas, cabecalhos, largura_total_cm=16.5):
    n = len(colunas)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(cabecalhos):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for linha in linhas:
        row = table.add_row().cells
        for i, c in enumerate(colunas):
            row[i].text = str(linha.get(c, ""))
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def to_float(v):
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


def fmt(v, casas=4):
    f = to_float(v)
    return f"{f:.{casas}f}" if f is not None else "-"


def add_image_centered(doc, caminho, largura_cm=15.5):
    if caminho.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(caminho), width=Cm(largura_cm))


def main():
    linhas = carregar_csv()
    dataset = json.loads((AVAL / "dataset.json").read_text(encoding="utf-8"))
    grafo_comp = carregar_json("fase7_comparacao_grafo.json") or []

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    set_base_style(doc)

    # 1. Capa
    add_title_page(doc)

    # 2. Fonte e dataset
    add_heading(doc, "1. A fonte e o dataset de avaliação")
    tipos = {}
    for q in dataset["queries_benchmark"]:
        tipos[q["tipo"]] = tipos.get(q["tipo"], 0) + 1
    doc.add_paragraph(
        f"O corpus de gabarito ('documentos') foi construído no nível de artigo: "
        f"{len(dataset['documentos'])} artigos únicos (incluindo sufixos de emenda "
        f"como Art. 149-A a 149-G), extraídos via regex sobre o texto completo da "
        f"Constituição (excluindo o ADCT, que reinicia a numeração de artigos e "
        f"causaria colisão de IDs)."
    )
    doc.add_paragraph(
        f"O gabarito de perguntas (avaliacao/dataset.json) tem {len(dataset['queries_benchmark'])} "
        f"perguntas, divididas em 3 tipos: {tipos.get('factual',0)} factuais (respondidas por 1 "
        f"artigo), {tipos.get('reformulavel',0)} reformuláveis (vocabulário coloquial, boas para "
        f"testar reescrita de consulta) e {tipos.get('multi_hop',0)} temáticas/multi-hop (2-3 "
        f"artigos, boas para RAG-Fusion e grafo). Todas as perguntas foram escritas manualmente "
        f"em linguagem coloquial — sem copiar termos exatos do texto legal — e verificadas contra "
        f"o conteúdo real dos artigos extraídos, para evitar o viés de 'pergunta que já bate com "
        f"o vocabulário do documento'."
    )
    doc.add_paragraph(
        "Matching pergunta→artigo (para Hit@K/Recall@K/MRR/NDCG@K): o campo 'id_original' "
        "referenciado em app/consulta.py nunca é populado em nenhum lugar do projeto original "
        "(cada chunk só recebe meta={\"arquivo\": nome_do_arquivo}). Como a fonte é um único PDF, "
        "isso não serve de granularidade de gabarito. A solução adotada foi identificar quais "
        "artigos um chunk recuperado cobre via regex sobre o próprio conteúdo do chunk "
        "(Art\\.?\\s*\\d+[º°]?(-[A-Z])?) — funciona para qualquer técnica de chunking, sem exigir "
        "nenhuma mudança no pipeline de indexação."
    )

    # 3. Metodologia
    add_heading(doc, "2. Metodologia")
    doc.add_paragraph(
        "Métricas de recuperação: Hit@5, Recall@5, MRR (calculado sobre toda a lista "
        "recuperada) e NDCG@10 (graduado: nota 2 = muito relevante, nota 1 = relevante), "
        "implementadas em avaliacao/comum.py (NDCG via sklearn.metrics.ndcg_score)."
    )
    doc.add_paragraph(
        "Para não gastar uma chamada de LLM por pergunta apenas para medir recuperação, "
        "app/busca_avancada.construir() ganhou um parâmetro novo apenas_recuperacao=True, "
        "que monta o mesmo pipeline Haystack de cada técnica só até a etapa de busca, "
        "sem os nós de prompt/geração final. As técnicas de reescrita de consulta "
        "(multi_query/rag_fusion/step_back) continuam fazendo sua própria chamada de LLM "
        "de reescrita — isso é parte da técnica, não da geração da resposta."
    )
    doc.add_paragraph(
        "Cada experimento reindexa o OpenSearch do zero (avaliacao/reindexar.py apaga o "
        "índice antes de indexar) para garantir comparação controlada — uma variável por "
        "experimento, mesmo dataset, mesmo top_k quando não é a variável testada."
    )
    doc.add_paragraph(
        "Ferramentas/padrões reaproveitados sem modificar: bench_embeddings/app "
        "(comparação de modelos de embedding, Fase 3), o padrão de busca híbrida "
        "BM25+RRF da Aula 4 (OpenSearchHybridRetriever), o padrão de reranking da Aula 3 "
        "(TransformersSimilarityRanker/BAAI/bge-reranker-v2-m3) e o padrão RAGAS+Groq das "
        "Aulas 5/8 (ResponseRelevancy(strictness=1) — a Groq só aceita n=1)."
    )
    doc.add_paragraph(
        "Observabilidade (LangFuse): a primeira tentativa de subir o LangFuse via o "
        "docker-compose do curso falhou (a imagem langfuse:latest já é a v3, que exige "
        "ClickHouse — não presente naquele compose, só Postgres). A solução foi clonar o "
        "repositório oficial do LangFuse (github.com/langfuse/langfuse), que já traz o "
        "compose completo da v3 (Postgres + ClickHouse + Redis + MinIO + web + worker), e "
        "subir via Podman. Com um projeto e chaves novos gerados na UI, o tracing "
        "funcionou de ponta a ponta: uma consulta de teste com a melhor configuração "
        "(rerank) gerou um trace completo, com latência por componente — embedder 4.59s, "
        "retriever 0.07s, ranker (cross-encoder) 16.43s, llm 1.94s (2 222 tokens de prompt "
        "→ 245 de completion) — confirmando, com dados reais de produção, o mesmo "
        "diagnóstico de custo/latência do reranking já visto nas métricas offline da "
        "Fase 6 (Seção 4.6): o cross-encoder é, de longe, o componente mais caro do "
        "pipeline."
    )

    # 4. Baseline
    add_heading(doc, "3. Baseline (Fase 0)")
    doc.add_paragraph(
        "Rodado 'como veio': chunking=auto, estratégia=opensearch, embedding=nomic-embed-text, "
        "técnica=baseline, top_k=5."
    )
    l0 = next(l for l in linhas if l["exp"] == "exp01")
    doc.add_paragraph(
        f"Resultado: Hit@5={l0['hit@5']}, Recall@5={l0['recall@5']}, MRR={l0['mrr']}, "
        f"NDCG@10={l0['ndcg@10']}."
    ).runs[0].font.bold = True
    doc.add_paragraph(
        "Achado central: o agente de extração escolheu 'extrair_texto' (Docling), que "
        "falhou silenciosamente a partir da página 82 de 164 (std::bad_alloc no modelo de "
        "layout, repetido em todas as páginas seguintes) — sem lançar exceção que "
        "acionasse o fallback para PyMuPDF já previsto em app/extracao.py (o limiar de "
        "fallback é <50 caracteres extraídos; um truncamento parcial de ~80% do texto não "
        "é pego por esse limiar). Resultado: só 127 mil dos 698 mil caracteres do PDF "
        "foram extraídos (~18%), faltando a maior parte dos artigos no índice."
    )

    # 5. Experimentos (fases 1-8)
    add_heading(doc, "4. Experimentos")

    add_heading(doc, "4.1 Fase 1 — Extração", nivel=2)
    doc.add_paragraph(
        "Hipótese: texto mal extraído limita tudo a jusante. "
        "Mudança: avaliacao/reindexar.py --fonte pymupdf substitui a extração por "
        "PyMuPDF completo (698 372 caracteres, 100% dos 250 artigos)."
    )
    subset1 = [l for l in linhas if l["exp"] in ("exp01", "exp02b", "exp02")]
    add_metric_table(doc, subset1, ["exp", "mudanca", "hit@5", "recall@5", "mrr", "ndcg@10"],
                      ["exp", "Configuração", "Hit@5", "Recall@5", "MRR", "NDCG@10"])
    doc.add_paragraph(
        "\nAnálise: isolando só a extração (exp02b vs exp01, mesmo chunking hierárquico "
        "nos dois), todas as métricas praticamente dobram ou triplicam. Este é o achado "
        "mais forte de toda a jornada — um bug de extração (não uma questão de qualidade "
        "fina) derrubou o Recall em cerca de 80%, e o próprio pipeline não teve como "
        "perceber sozinho, pois a falha por página é silenciosa e não aciona o fallback "
        "existente. Isso confirma, de forma extrema, a regra de ouro do roteiro: medir "
        "antes de otimizar."
    )

    add_heading(doc, "4.2 Fase 2 — Chunking", nivel=2)
    doc.add_paragraph(
        "Mudança: chunking forçado nas 5 técnicas nativas de app/indexacao.py::chunkar, "
        "com extração (PyMuPDF) e embedding (nomic-embed-text) fixos."
    )
    subset2 = [l for l in linhas if l["fase"] == "Fase 2 - Chunking"]
    subset2_sorted = sorted(subset2, key=lambda l: -to_float(l["ndcg@10"]))
    add_metric_table(doc, subset2_sorted, ["exp", "hit@5", "recall@5", "mrr", "ndcg@10"],
                      ["Técnica", "Hit@5", "Recall@5", "MRR", "NDCG@10"])
    doc.add_paragraph(
        "\nAnálise: contra a intuição de que técnicas mais sofisticadas (semântica, "
        "hierárquica) ganhariam, quem venceu foi o chunking fixo (200 palavras), com "
        "recursivo logo atrás. Hipótese explicativa: artigos da Constituição variam "
        "muito de tamanho — o Art. 5º sozinho tem 78 incisos e ~14 mil caracteres — "
        "então janelas fixas de 200 palavras cortam o texto em unidades mais homogêneas "
        "e densas por chunk, enquanto sentença_janela fragmenta demais um artigo longo "
        "(muitos chunks quase-duplicados) e hierárquico/semântico às vezes agrupam vários "
        "artigos curtos num mesmo chunk, diluindo a especificidade do embedding."
    )

    add_heading(doc, "4.3 Fase 3 — Modelo de embedding", nivel=2)
    doc.add_paragraph(
        "Bancada bench_embeddings (274 artigos x 22 perguntas, k=10), comparando "
        "nomic-embed-text (768d, padrão), bge-m3 (1024d) e mxbai-embed-large (1024d):"
    )
    bench = carregar_json("../fase3_resultado_embeddings.json") or carregar_json("fase3_resultado_embeddings.json")
    if bench:
        add_metric_table(
            doc, bench["resultados"],
            ["modelo", "hit@10", "recall@10", "mrr", "ndcg@10", "auc", "latencia_s"],
            ["Modelo", "Hit@10", "Recall@10", "MRR", "NDCG@10", "AUC", "Latência (s)"],
        )
    l3 = next(l for l in linhas if l["exp"] == "exp03")
    doc.add_paragraph(
        f"\nbge-m3 venceu disparado na bancada (NDCG@10=0.9215 vs 0.53/0.54). Reindexando "
        f"o app com chunking=fixo + bge-m3: Hit@5={l3['hit@5']}, Recall@5={l3['recall@5']}, "
        f"MRR={l3['mrr']}, NDCG@10={l3['ndcg@10']} — quase o dobro do melhor resultado da "
        f"Fase 2. mxbai-embed-large tem janela de contexto de 512 tokens (bem menor que "
        f"bge-m3/nomic, 8192) — pode truncar artigos longos como o Art. 5º, o que ajuda a "
        f"explicar seu desempenho fraco apesar da mesma dimensão de bge-m3."
    )

    add_heading(doc, "4.4 Fase 4 — Recuperação base (top_k e busca híbrida)", nivel=2)
    subset4k = [l for l in linhas if "topk" in l["exp"] and "real" in l["exp"]]
    doc.add_paragraph("top_k real (Hit@k/Recall@k/NDCG@k calculado no próprio k, chunking=fixo+bge-m3):")
    add_metric_table(doc, subset4k, ["exp", "hit@5", "recall@5", "mrr", "ndcg@10"],
                      ["top_k", "Hit@k", "Recall@k", "MRR", "NDCG@k"])
    doc.add_paragraph(
        "\nRecall sobe consistentemente com top_k (mais chunks candidatos = mais chance de "
        "cobrir o artigo certo), com custo de mais contexto (latência/tokens) na geração."
    )
    l4h = next(l for l in linhas if l["exp"] == "exp04_hibrida")
    doc.add_paragraph(
        f"Nova técnica 'hibrida' (BM25 + denso, fundidos por RRF via "
        f"OpenSearchHybridRetriever, implementada em app/busca_avancada.py), top_k=5: "
        f"Hit@5={l4h['hit@5']}, Recall@5={l4h['recall@5']}, MRR={l4h['mrr']}, "
        f"NDCG@10={l4h['ndcg@10']} — melhora sobre a busca densa pura, a melhor "
        f"configuração de recuperação até este ponto da jornada."
    )

    add_heading(doc, "4.5 Fase 5 — Query enhancement", nivel=2)
    subset5 = [l for l in linhas if l["fase"] == "Fase 5 - Query enhancement"] + [l4h] + [l3]
    subset5_sorted = sorted(subset5, key=lambda l: -to_float(l["ndcg@10"]))
    add_metric_table(doc, subset5_sorted, ["exp", "hit@5", "recall@5", "mrr", "ndcg@10"],
                      ["Técnica", "Hit@5", "Recall@5", "MRR", "NDCG@10"])
    doc.add_paragraph(
        "\nAnálise: nenhuma técnica de reescrita de consulta (multi_query, rag_fusion, "
        "step_back) supera a busca híbrida, e nenhuma supera sequer a busca densa pura "
        "com bge-m3. Quando o embedding de base já é forte e as perguntas do gabarito, "
        "embora coloquiais, são objetivas, reescrever a consulta introduz variações que "
        "dispersam o ranking (o MRR cai em todas as 3 técnicas). O ganho de multi-query/ "
        "RAG-Fusion tende a aparecer quando o embedding de base é fraco ou a pergunta é "
        "muito ambígua — não é o caso deste corpus com bge-m3."
    )

    add_heading(doc, "4.6 Fase 6 — Reranking", nivel=2)
    l6 = next(l for l in linhas if l["exp"] == "exp06_rerank")
    doc.add_paragraph(
        "Nova técnica 'rerank' (app/busca_avancada.py): recupera os top-20 candidatos "
        "densos e reordena com um cross-encoder (BAAI/bge-reranker-v2-m3, via "
        "TransformersSimilarityRanker) para o top-5 final."
    )
    add_metric_table(
        doc, [l3, l4h, l6], ["exp", "hit@5", "recall@5", "mrr", "ndcg@10", "latencia_s"],
        ["Config", "Hit@5", "Recall@5", "MRR", "NDCG@10", "Latência méd. (s)"],
    )
    doc.add_paragraph(
        f"\nMaior salto de toda a recuperação: MRR sobe de {l4h['mrr']} (híbrida) para "
        f"{l6['mrr']}, e NDCG@10 de {l4h['ndcg@10']} para {l6['ndcg@10']} — exatamente o "
        f"padrão previsto no roteiro ('MRR e NDCG são os que mais devem subir'). O custo "
        f"é latência: {l6['latencia_s']}s/pergunta em CPU (vs ~2.7s da busca densa), "
        f"porque o cross-encoder reavalia 20 pares (pergunta, chunk) por consulta. "
        f"Esta é a melhor configuração de recuperação de toda a jornada."
    )

    add_heading(doc, "4.7 Fase 7 — Avançada: Grafo (LightRAG)", nivel=2)
    doc.add_paragraph(
        "Escopo: em vez da Constituição inteira (o que geraria centenas de chamadas de "
        "LLM para extração de entidades/relações via Groq — testado e confirmado: "
        "estourou tanto o limite por minuto quanto o limite diário de tokens do tier "
        "gratuito da Groq), o grafo foi construído só com o Título II — Dos Direitos e "
        "Garantias Fundamentais (Art. 5º a 17), um subconjunto coeso e rico em remissão "
        "interna, bom para perguntas multi-hop. O grafo resultante tem 152 nós e 99 "
        "arestas (entidades como Artigo 5, Judiciário, Executivo, Cidadão, Liberdade, "
        "Igualdade, Propriedade e as relações entre elas)."
    )
    doc.add_paragraph(
        "Limitação de instrumentação: consulta.consultar_grafo() devolve fontes = "
        "[{\"id\": \"grafo\", ...}] (sem IDs reais de artigo), então não dá para calcular "
        "Hit@K/Recall@K/NDCG@K no grafo com a mesma metodologia do OpenSearch. A "
        "comparação foi feita de forma qualitativa, nas 3 perguntas multi-hop do "
        "gabarito com sobreposição no Título II (Q16, Q18, Q21)."
    )
    for item in grafo_comp:
        doc.add_paragraph(f"\n{item['id']}: {item['query']}").runs[0].font.bold = True
        doc.add_paragraph(f"Referência: {item['resposta_referencia']}")
        p_os = doc.add_paragraph()
        p_os.add_run("OpenSearch (índice completo): ").italic = True
        p_os.add_run(item["opensearch"][:500] + ("..." if len(item["opensearch"]) > 500 else ""))
        p_gr = doc.add_paragraph()
        p_gr.add_run("Grafo (Título II): ").italic = True
        p_gr.add_run(item["grafo"][:500] + ("..." if len(item["grafo"]) > 500 else ""))
    doc.add_paragraph(
        "\nAnálise: nas perguntas totalmente cobertas pelo Título II (Q16, Q18), o grafo "
        "produz respostas bem sintetizadas e corretas, mas com uma tendência a "
        "'enriquecer' a resposta com leis, decretos e programas específicos (Lei do SUS, "
        "Bolsa Família, escolas cívico-militares) que NÃO estavam no texto ingerido — "
        "conhecimento de mundo do LLM sendo misturado à síntese do grafo, um risco de "
        "faithfulness. Em Q21, que depende de um artigo fora do escopo ingerido (Art. "
        "227, sobre crianças), o efeito fica mais grave: o OpenSearch (que tem o índice "
        "completo) cita corretamente o Art. 227 com seus incisos; o grafo, sem esse "
        "artigo na sua base, atribui as garantias de criança a 'Artigo 5, 6, 7' — uma "
        "atribuição incorreta, mas apresentada com a mesma confiança. O próprio texto "
        "gerado pelo grafo chega a admitir a limitação ('a Reference Document List está "
        "vazia'), o que sugere que o grafo tem consciência de faltar contexto, mas ainda "
        "assim gera uma resposta com aparência de completude. Isso reforça: um grafo "
        "construído sobre um subconjunto do corpus é útil para perguntas multi-hop DENTRO "
        "do seu escopo, mas arrisca respostas confiantes e incorretas fora dele — "
        "diferente da recuperação lexical/densa, que simplesmente não traria contexto "
        "sobre o que não foi indexado."
    )

    add_heading(doc, "4.8 Fase 8 — RAGAS (geração)", nivel=2)
    m_base = carregar_json("../ragas_exp08_baseline_medias.json") or carregar_json("ragas_exp08_baseline_medias.json")
    m_rer = carregar_json("../ragas_exp08_rerank_medias.json") or carregar_json("ragas_exp08_rerank_medias.json")
    doc.add_paragraph(
        "RAGAS (Faithfulness, ResponseRelevancy(strictness=1), LLMContextRecall, "
        "LLMContextPrecisionWithReference; juiz Groq llama-3.3-70b-versatile, embeddings "
        "Ollama nomic-embed-text) rodado no pipeline completo (busca + geração) para a "
        "técnica baseline (densa) e para rerank (melhor configuração), ambos sobre o "
        "índice chunking=fixo + bge-m3, 22 perguntas:"
    )
    if m_base and m_rer:
        linhas_ragas = [
            {"config": "baseline (densa)", **m_base},
            {"config": "rerank (melhor config)", **m_rer},
        ]
        add_metric_table(
            doc, linhas_ragas,
            ["config", "faithfulness", "answer_relevancy", "context_recall", "llm_context_precision_with_reference"],
            ["Config", "Faithfulness", "Answer Relevancy", "Context Recall", "Context Precision (ref)"],
        )
    doc.add_paragraph(
        "\nAnálise: rerank melhora Faithfulness (+2.2 p.p.) e Answer Relevancy (+2.2 p.p.) "
        "em relação à busca densa pura — reordenar os candidatos por relevância real "
        "(cross-encoder) entrega ao LLM um contexto mais focado, gerando respostas mais "
        "fiéis e mais diretamente relevantes à pergunta. Curiosamente, Context Recall e "
        "Context Precision (with reference) caem ligeiramente (~0.8 p.p. cada) — o "
        "reranking reordena os MESMOS 20 candidatos, então às vezes troca um chunk "
        "'tecnicamente relevante' por outro mais bem escrito/focado, sem necessariamente "
        "aumentar a cobertura bruta do contexto. Isso é um trade-off real: qualidade "
        "percebida da resposta (faithfulness/relevancy) sobe mais do que a cobertura "
        "bruta do contexto (recall/precision) desce — vale a pena para este corpus, mas "
        "é o tipo de nuance que só aparece medindo os dois lados (recuperação e geração) "
        "separadamente, como pede o roteiro."
    )

    # 6. Tabela consolidada + gráficos
    doc.add_page_break()
    add_heading(doc, "5. Tabela consolidada e gráficos")
    add_metric_table(
        doc, linhas, ["exp", "fase", "hit@5", "recall@5", "mrr", "ndcg@10"],
        ["exp", "Fase", "Hit@5", "Recall@5", "MRR", "NDCG@10"],
    )
    doc.add_paragraph()
    add_image_centered(doc, GRAF / "evolucao_ndcg_at_10.png")
    doc.add_paragraph("Evolução do NDCG@10 ao longo da jornada de experimentos.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_centered(doc, GRAF / "evolucao_recall_at_5.png")
    doc.add_paragraph("Evolução do Recall@5.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_centered(doc, GRAF / "fase2_chunking.png")
    doc.add_paragraph("Comparação das 5 técnicas de chunking (Fase 2).").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_centered(doc, GRAF / "fase5_query_enhancement.png")
    doc.add_paragraph("Comparação das técnicas de query enhancement (Fase 5).").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_image_centered(doc, GRAF / "ragas_comparativo.png")
    doc.add_paragraph("RAGAS: baseline vs melhor configuração (rerank).").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 7. Melhor config final
    doc.add_page_break()
    add_heading(doc, "6. Melhor configuração final")
    doc.add_paragraph(
        "Extração: PyMuPDF (fallback nativo do projeto, usado no lugar do Docling, que "
        "falha por falta de memória neste documento).\n"
        "Chunking: fixo (200 palavras, sem sobreposição).\n"
        "Embedding: bge-m3 (Ollama, 1024 dimensões).\n"
        "Recuperação: densa top-20 → reranking com BAAI/bge-reranker-v2-m3 → top-5.\n"
        "top_k final para geração: 5."
    )
    l_base = l0
    delta = {
        "Hit@5": (to_float(l6["hit@5"]) - to_float(l_base["hit@5"])),
        "Recall@5": (to_float(l6["recall@5"]) - to_float(l_base["recall@5"])),
        "MRR": (to_float(l6["mrr"]) - to_float(l_base["mrr"])),
        "NDCG@10": (to_float(l6["ndcg@10"]) - to_float(l_base["ndcg@10"])),
    }
    doc.add_paragraph(
        f"\nGanho total vs baseline (Fase 0): Hit@5 {l_base['hit@5']} → {l6['hit@5']} "
        f"(Δ+{delta['Hit@5']:.4f}); Recall@5 {l_base['recall@5']} → {l6['recall@5']} "
        f"(Δ+{delta['Recall@5']:.4f}); MRR {l_base['mrr']} → {l6['mrr']} "
        f"(Δ+{delta['MRR']:.4f}); NDCG@10 {l_base['ndcg@10']} → {l6['ndcg@10']} "
        f"(Δ+{delta['NDCG@10']:.4f})."
    ).runs[0].font.bold = True
    doc.add_paragraph(
        "A extração (Fase 1) e a troca de embedding (Fase 3) foram, isoladamente, os dois "
        "maiores saltos da jornada; chunking e reranking foram os ajustes finos que "
        "levaram a configuração ao topo."
    )

    # 8. Análise crítica
    add_heading(doc, "7. Análise crítica")
    add_heading(doc, "Trade-offs (qualidade × latência × custo)", nivel=2)
    doc.add_paragraph(
        "• Reranking: melhor qualidade de ranking (MRR/NDCG) e melhor Faithfulness/"
        "Answer Relevancy, mas ~6x mais latência que a busca densa pura (16s vs 2.7s por "
        "pergunta, rodando o cross-encoder em CPU) — em produção, valeria avaliar GPU ou "
        "um cross-encoder menor.\n"
        "• Busca híbrida (BM25+RRF): ganho modesto sobre a densa pura, praticamente sem "
        "custo adicional de latência — bom custo-benefício quando não se quer pagar o "
        "preço do reranking.\n"
        "• Grafo (LightRAG): o mais caro de todos em tokens de LLM — mesmo um subconjunto "
        "pequeno (13 artigos, ~35 chunks) consumiu boa parte da cota diária gratuita da "
        "Groq (100 000 tokens/dia) só na fase de indexação (extração de entidades e "
        "relações), além de exigir reduzir a concorrência do LightRAG (max_async=1) para "
        "não estourar o limite por minuto (12 000 TPM). Indexar a Constituição inteira no "
        "grafo seria proibitivo num tier gratuito."
    )
    add_heading(doc, "O que NÃO funcionou e por quê", nivel=2)
    doc.add_paragraph(
        "• Query enhancement (multi_query, rag_fusion, step_back) não melhorou nada sobre "
        "a busca densa/híbrida neste corpus — o embedding bge-m3 já é forte o suficiente "
        "para as perguntas coloquiais do gabarito, e reescrever a consulta introduziu mais "
        "ruído do que ganho de cobertura.\n"
        "• Chunking 'sofisticado' (semântico, hierárquico) perdeu para o chunking fixo — "
        "a estrutura muito irregular dos artigos constitucionais (de 1 linha a 14 mil "
        "caracteres) não se beneficia das heurísticas de agrupamento por tópico/estrutura "
        "do jeito que se beneficiaria um texto mais homogêneo.\n"
        "• O grafo, fora do seu escopo de ingestão, gerou respostas confiantes mas "
        "incorretas (atribuindo direitos do Art. 227 aos Art. 5/6/7) — um risco real de "
        "uso de grafos parciais em produção."
    )
    add_heading(doc, "Limitações do dataset e da metodologia", nivel=2)
    doc.add_paragraph(
        "• 22 perguntas é uma amostra pequena; os deltas absolutos entre técnicas próximas "
        "(ex.: fixo vs recursivo na Fase 2) podem não ser estatisticamente robustos.\n"
        "• O matching pergunta→artigo via regex sobre o conteúdo do chunk é uma aproximação: "
        "chunks grandes que mencionam um número de artigo de passagem (ex.: uma remissão "
        "'nos termos do art. X') podem contar como 'contendo' aquele artigo sem "
        "efetivamente tratar do seu conteúdo.\n"
        "• A comparação do grafo (Fase 7) não usou as mesmas métricas de ranking do "
        "restante da jornada, por limitação do próprio projeto (fontes do grafo não "
        "carregam IDs) — a comparação ficou qualitativa.\n"
        "• O trace do LangFuse (Seção 2) cobre só uma consulta de verificação pontual, "
        "não as 22 perguntas do gabarito nem todas as técnicas — serviu para validar a "
        "instrumentação e confirmar o custo do reranking, não como fonte sistemática de "
        "métricas desta jornada."
    )

    # 9. Conclusão
    add_heading(doc, "8. Conclusão")
    doc.add_paragraph(
        "A jornada partiu de uma baseline severamente prejudicada por um problema de "
        "extração não relacionado a nenhuma técnica de RAG em si (Docling ficando sem "
        "memória a partir da metade do PDF) — o que reforça a lição central do roteiro: "
        "meça antes de otimizar, porque às vezes o maior ganho não está em nenhuma "
        "técnica de busca, e sim em garantir que o dado que entra no pipeline está "
        "completo. A partir daí, embedding (bge-m3) e reranking (bge-reranker-v2-m3) "
        "foram as duas alavancas de maior impacto; chunking teve impacto moderado e "
        "favoreceu, contra a intuição, a técnica mais simples; e técnicas de reescrita "
        "de consulta não ajudaram neste corpus específico. O grafo de conhecimento "
        "mostrou-se poderoso para sintetizar respostas dentro do seu escopo de ingestão, "
        "mas caro (em tokens de LLM) e arriscado (respostas confiantes fora do escopo) "
        "quando indexado parcialmente — uma decisão de produto que precisaria pesar "
        "custo, cobertura e o risco de alucinação por escopo incompleto."
    )
    doc.add_paragraph(
        "Próximos passos: (1) rodar o gabarito completo (22 perguntas x técnicas) com o "
        "LangFuse já instrumentado, usando os traces para diagnosticar casos individuais "
        "de falha; (2) testar o grafo sobre a Constituição inteira usando um provedor de "
        "LLM com cota maior; (3) ampliar o gabarito para mais perguntas por tipo, "
        "aumentando a robustez estatística das comparações; (4) avaliar reranking em GPU "
        "para viabilizar seu uso em produção com latência aceitável."
    )

    # 10. Anexos
    doc.add_page_break()
    add_heading(doc, "9. Anexos")
    doc.add_paragraph(
        "Anexo A — Log completo do erro de extração do Docling (std::bad_alloc, páginas "
        "82-164). Anexo B — comparações completas de todas as respostas do grafo x "
        "OpenSearch (Q16/Q18/Q21) reproduzidas na Seção 4.7. Anexo C — código-fonte "
        "completo em aula12/projeto_final/avaliacao/ e as duas técnicas novas em "
        "app/busca_avancada.py (hibrida, rerank). Anexo D — trace do LangFuse da consulta "
        "de verificação (melhor configuração, técnica=rerank), navegável em "
        "http://localhost:3000/project/cmrbkmy130006rs02n5hqj73k/traces/"
        "155d8617f1a861d30b6209103be50996 (instância local); breakdown de latência por "
        "componente: embedder 4.59s, retriever 0.07s, ranker (cross-encoder) 16.43s, "
        "llm 1.94s (2 222 tokens de prompt → 245 de completion)."
    )

    out = PASTA / "Relatorio_Final_RAG_Constituicao.docx"
    doc.save(str(out))
    print(f"Relatorio salvo em {out}")


if __name__ == "__main__":
    main()
